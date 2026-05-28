"""
Convert every live legal .jsx in web/src/pages/legal/ to a clean .docx in docs/legal/.

The .jsx files in web/src/pages/legal/ are the source of truth (they're what
ships to myrxfit.com). The .docx files in docs/legal/ are read-only snapshots
for offline review — open them in Word, mark up edits, hand to lawyers.

JSX -> DOCX mapping:
  <LegalLayout title=X effectiveDate=Y> -> Title page (Title + italic effective date)
  <h1>/<h2>/<h3>                       -> Heading 1/2/3
  <p>                                  -> Normal paragraph
  <ul><li>                             -> List Bullet
  <ol><li>                             -> List Number
  <strong>/<b>                         -> bold run
  <em>/<i>                             -> italic run
  <a href="...">text</a>               -> real Word hyperlink (blue + underlined)
  {' '} expressions                    -> single space (JSX whitespace token)
  any other {...} expression           -> stripped (legal docs are static prose)

Run from repo root:
    python scripts/build_legal_docx.py
"""
from __future__ import annotations
import re
import html.parser
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
import docx.opc.constants
import docx.oxml.shared

ROOT = Path(__file__).resolve().parents[1]
SRC_DIR = ROOT / "web" / "src" / "pages" / "legal"   # source of truth (the live React components)
OUT_DIR = ROOT / "docs" / "legal"                    # where .docx archives land

DOCS = [
    "AcceptableUsePolicy.jsx",
    "CoachAgreement.jsx",
    "CookiePolicy.jsx",
    "DataProcessingAgreement.jsx",
    "HealthDisclaimer.jsx",
    "HowWeCompute.jsx",
    "PrivacyPolicy.jsx",
    "RefundPolicy.jsx",
    "TermsOfService.jsx",
]

# ─────────── helpers ────────────────────────────────────────────────────────────
def extract_legal_layout(jsx_text: str):
    """Pull the LegalLayout title + effectiveDate + inner JSX."""
    m = re.search(r"<LegalLayout\s+([^>]*)>(.*?)</LegalLayout>", jsx_text, re.DOTALL)
    if not m:
        return None, None, None
    attrs, inner = m.group(1), m.group(2)
    title_m = re.search(r'title="([^"]+)"', attrs)
    date_m = re.search(r'effectiveDate="([^"]+)"', attrs)
    return (
        title_m.group(1) if title_m else "Untitled",
        date_m.group(1) if date_m else None,
        inner,
    )

def clean_jsx(inner: str) -> str:
    """Turn JSX-ish content into HTML-parseable string.

    - `{' '}` and similar whitespace-only expressions -> ' '
    - Any other `{...}` expression -> stripped (static legal text only)
    - Self-closing comments and stray newlines collapsed
    """
    # whitespace JSX tokens like {' '}, {`\n`}, etc.
    inner = re.sub(r"\{\s*['\"`]\s*['\"`]\s*\}", " ", inner)
    # Drop any other JSX expressions ({var}, {2026}, etc.)
    inner = re.sub(r"\{[^{}]*\}", "", inner)
    # Strip JSX-style HTML comments
    inner = re.sub(r"\{/\*.*?\*/\}", "", inner, flags=re.DOTALL)
    # Wrap in a root so html.parser is happy with multiple top-level siblings
    return f"<root>{inner}</root>"

def add_hyperlink(paragraph, url: str, text: str) -> None:
    """python-docx has no native hyperlink API — drop down to raw OXML."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(docx.oxml.shared.qn("r:id"), r_id)

    new_run = docx.oxml.shared.OxmlElement("w:r")
    rPr = docx.oxml.shared.OxmlElement("w:rPr")
    color = docx.oxml.shared.OxmlElement("w:color")
    color.set(docx.oxml.shared.qn("w:val"), "0563C1")  # Word link blue
    rPr.append(color)
    u = docx.oxml.shared.OxmlElement("w:u")
    u.set(docx.oxml.shared.qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)

    text_elem = docx.oxml.shared.OxmlElement("w:t")
    text_elem.text = text
    text_elem.set(docx.oxml.shared.qn("xml:space"), "preserve")
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

# ─────────── renderer ───────────────────────────────────────────────────────────
class DocxRenderer(html.parser.HTMLParser):
    def __init__(self, doc):
        super().__init__(convert_charrefs=True)
        self.doc = doc
        self.current_paragraph = None
        self.list_kind = None        # 'ul' or 'ol' (None if outside list)
        self.in_link = False
        self.link_url = None
        self.formatting = []         # stack of 'bold' / 'italic' currently open

    def _ensure_paragraph(self):
        """If a stray text node arrives outside any block tag, start a Normal paragraph."""
        if self.current_paragraph is None:
            self.current_paragraph = self.doc.add_paragraph(style="Normal")

    def handle_starttag(self, tag, attrs):
        attrs_d = dict(attrs)
        if tag in ("h1", "h2", "h3"):
            level = int(tag[1])
            self.current_paragraph = self.doc.add_heading(level=level)
        elif tag == "p":
            self.current_paragraph = self.doc.add_paragraph(style="Normal")
        elif tag in ("ul", "ol"):
            self.list_kind = tag
        elif tag == "li":
            style = "List Bullet" if self.list_kind == "ul" else "List Number"
            self.current_paragraph = self.doc.add_paragraph(style=style)
        elif tag in ("strong", "b"):
            self.formatting.append("bold")
        elif tag in ("em", "i"):
            self.formatting.append("italic")
        elif tag == "a":
            self.in_link = True
            self.link_url = attrs_d.get("href", "")
        elif tag == "br":
            if self.current_paragraph:
                self.current_paragraph.add_run("\n")
        # silently ignore other tags (root, div wrappers, etc.)

    def handle_endtag(self, tag):
        if tag in ("h1", "h2", "h3", "p", "li"):
            self.current_paragraph = None
        elif tag in ("ul", "ol"):
            self.list_kind = None
        elif tag in ("strong", "b") and "bold" in self.formatting:
            self.formatting.remove("bold")
        elif tag in ("em", "i") and "italic" in self.formatting:
            self.formatting.remove("italic")
        elif tag == "a":
            self.in_link = False
            self.link_url = None

    def handle_data(self, data):
        # Collapse whitespace runs (HTML semantics)
        text = re.sub(r"\s+", " ", data)
        if not text.strip():
            # Preserve a single space between adjacent inline runs
            if self.current_paragraph and self.current_paragraph.text and not self.current_paragraph.text.endswith(" "):
                self.current_paragraph.add_run(" ")
            return
        self._ensure_paragraph()
        if self.in_link and self.link_url:
            add_hyperlink(self.current_paragraph, self.link_url, text)
        else:
            run = self.current_paragraph.add_run(text)
            if "bold" in self.formatting:
                run.bold = True
            if "italic" in self.formatting:
                run.italic = True

# ─────────── main loop ──────────────────────────────────────────────────────────
def convert_one(jsx_path: Path) -> None:
    docx_path = OUT_DIR / (jsx_path.stem + ".docx")
    text = jsx_path.read_text(encoding="utf-8")
    title, effective_date, inner = extract_legal_layout(text)
    if title is None:
        print(f"[WARN] {jsx_path.name} has no <LegalLayout>; skipping.")
        return

    doc = Document()
    # Title page header
    title_p = doc.add_heading(title, level=0)
    if effective_date:
        date_p = doc.add_paragraph(f"Effective: {effective_date}")
        for r in date_p.runs:
            r.italic = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)  # slate-500
    doc.add_paragraph()  # spacer before body

    # Render body
    body_html = clean_jsx(inner)
    DocxRenderer(doc).feed(body_html)

    # Save (with PermissionError fallback if Word has it open)
    try:
        doc.save(docx_path)
        print(f"[OK] {jsx_path.name} -> {docx_path.name}")
    except PermissionError:
        alt = docx_path.with_name(docx_path.stem + "_v2.docx")
        doc.save(alt)
        print(f"[WARN] {docx_path.name} is open in Word — wrote {alt.name} instead.")

def main():
    if not SRC_DIR.exists():
        raise SystemExit(f"Missing source folder: {SRC_DIR}")
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for name in DOCS:
        jsx = SRC_DIR / name
        if not jsx.exists():
            print(f"[WARN] Source missing: {jsx}")
            continue
        convert_one(jsx)

if __name__ == "__main__":
    main()
